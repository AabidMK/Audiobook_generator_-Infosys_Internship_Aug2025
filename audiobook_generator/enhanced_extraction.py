import fitz  # PyMuPDF
import docx
import os
from datetime import datetime
import pytesseract
from PIL import Image
import io
import tempfile
from docx.document import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


class EnhancedTextExtraction:
    """
    Enhanced universal text extraction with OCR support for embedded images
    in PDF, DOCX, TXT, and standalone images. Combines text extraction with 
    Tesseract OCR for comprehensive content capture.
    """

    @staticmethod
    def configure_tesseract():
        """Configure Tesseract OCR settings for optimal performance"""
        # Uncomment and adjust path if needed for your system
        # pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'  # macOS Homebrew
        # pytesseract.pytesseract.tesseract_cmd = '/opt/homebrew/bin/tesseract'  # M1/M2 Mac
        
        # Optimized OCR configuration
        custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?;:" '
        return custom_config

    @staticmethod
    def extract_text_from_image(image_path_or_bytes, is_bytes=False):
        """
        Enhanced OCR extraction supporting both file paths and image bytes
        """
        supported_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}
        
        try:
            custom_config = EnhancedTextExtraction.configure_tesseract()
            
            if is_bytes:
                # Handle image bytes (from PDF/DOCX embedded images)
                img = Image.open(io.BytesIO(image_path_or_bytes))
            else:
                # Handle file path
                if not os.path.isfile(image_path_or_bytes):
                    return "Error: Image file not found"
                
                ext = os.path.splitext(image_path_or_bytes)[1].lower()
                if ext not in supported_formats:
                    return f"Error: Unsupported image format '{ext}'"
                
                img = Image.open(image_path_or_bytes)
            
            # Preprocessing for better OCR
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize if image is too small (improves OCR accuracy)
            width, height = img.size
            if width < 300 or height < 300:
                scale_factor = max(300/width, 300/height)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Extract text using Tesseract
            text = pytesseract.image_to_string(img, config=custom_config)
            
            # Clean up extracted text
            cleaned_text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
            
            return cleaned_text or ""
            
        except Exception as e:
            return f"Error extracting text from image: {e}"

    @staticmethod
    def extract_text_from_pdf_with_ocr(pdf_path):
        """
        Enhanced PDF extraction: regular text + OCR for images
        """
        if not os.path.isfile(pdf_path):
            return "Error: PDF file not found"
        
        try:
            extracted_text = []
            
            with fitz.open(pdf_path) as doc:
                for page_num, page in enumerate(doc):
                    print(f"Processing PDF page {page_num + 1}...")
                    
                    # Extract regular text
                    regular_text = page.get_text()
                    if regular_text.strip():
                        extracted_text.append(f"--- Page {page_num + 1} (Text) ---")
                        extracted_text.append(regular_text)
                    
                    # Extract and OCR images
                    image_list = page.get_images(full=True)
                    
                    for img_index, img in enumerate(image_list):
                        try:
                            # Get image data
                            xref = img[0]
                            pix = fitz.Pixmap(doc, xref)
                            
                            # Convert to PIL Image
                            if pix.n - pix.alpha < 4:  # GRAY or RGB
                                img_data = pix.tobytes("png")
                                ocr_text = EnhancedTextExtraction.extract_text_from_image(img_data, is_bytes=True)
                                
                                if ocr_text and not ocr_text.startswith("Error") and len(ocr_text.strip()) > 10:
                                    extracted_text.append(f"--- Page {page_num + 1} (Image {img_index + 1} OCR) ---")
                                    extracted_text.append(ocr_text)
                            
                            pix = None  # Release memory
                            
                        except Exception as img_error:
                            print(f"Warning: Could not process image {img_index + 1} on page {page_num + 1}: {img_error}")
                            continue
            
            result = '\n\n'.join(extracted_text).strip()
            return result or "Warning: PDF contains no extractable text or images"
            
        except Exception as e:
            return f"Error extracting PDF text: {e}"

    @staticmethod  
    def extract_images_from_docx(docx_path):
        """
        Extract embedded images from DOCX and perform OCR
        """
        try:
            doc = docx.Document(docx_path)
            ocr_texts = []
            
            # Get all relationships (including images)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        # Get image data
                        image_data = rel.target_part.blob
                        
                        # Perform OCR
                        ocr_text = EnhancedTextExtraction.extract_text_from_image(image_data, is_bytes=True)
                        
                        if ocr_text and not ocr_text.startswith("Error") and len(ocr_text.strip()) > 10:
                            ocr_texts.append(f"--- Embedded Image OCR ---")
                            ocr_texts.append(ocr_text)
                            
                    except Exception as img_error:
                        print(f"Warning: Could not process embedded image: {img_error}")
                        continue
            
            return '\n\n'.join(ocr_texts) if ocr_texts else ""
            
        except Exception as e:
            return f"Error extracting images from DOCX: {e}"

    @staticmethod
    def extract_text_from_docx_with_ocr(docx_path):
        """
        Enhanced DOCX extraction: regular text + OCR for embedded images
        """
        if not os.path.isfile(docx_path):
            return "Error: DOCX file not found"
        
        try:
            extracted_text = []
            
            # Extract regular text content
            doc = docx.Document(docx_path)
            
            # Extract paragraphs and tables
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text)
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            
            if texts:
                extracted_text.append("--- Document Text ---")
                extracted_text.append('\n'.join(texts))
            
            # Extract and OCR embedded images
            print("Extracting images from DOCX...")
            image_ocr_text = EnhancedTextExtraction.extract_images_from_docx(docx_path)
            
            if image_ocr_text:
                extracted_text.append(image_ocr_text)
            
            result = '\n\n'.join(extracted_text).strip()
            return result or "Warning: DOCX contains no extractable text or images"
            
        except Exception as e:
            return f"Error extracting DOCX text: {e}"

    @staticmethod
    def extract_text_from_txt(txt_path):
        """
        Same as before - TXT files don't contain images
        """
        if not os.path.isfile(txt_path):
            return "Error: TXT file not found"
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"Error extracting TXT text: {e}"
        return "Error: Unable to decode text file with UTF-8/16/latin-1/cp1252"

    @staticmethod
    def extract_text_from_any(file_path):
        """
        Universal extractor with OCR support
        """
        if not os.path.isfile(file_path):
            return f"Error: File '{file_path}' not found"
        
        ext = os.path.splitext(file_path)[1].lower()
        image_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}
        
        print(f"\n🔍 Processing: {os.path.basename(file_path)}")
        
        if ext == '.pdf':
            return EnhancedTextExtraction.extract_text_from_pdf_with_ocr(file_path)
        elif ext == '.docx':
            return EnhancedTextExtraction.extract_text_from_docx_with_ocr(file_path)
        elif ext == '.txt':
            return EnhancedTextExtraction.extract_text_from_txt(file_path)
        elif ext in image_formats:
            return EnhancedTextExtraction.extract_text_from_image(file_path)
        else:
            # Fallback for other text files
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                return "Error: Unknown file type and not UTF-8 decodable"
            except Exception as e:
                return f"Error reading unknown file type: {e}"

    @staticmethod
    def save_as_markdown(text, output_path, source_file=None):
        """
        Enhanced markdown saving with metadata
        """
        try:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            
            # Count different content types
            sections = text.split('---')
            text_sections = len([s for s in sections if 'Text' in s])
            image_sections = len([s for s in sections if 'Image' in s or 'OCR' in s])
            
            header = f"""# Extracted Content Report

**Source**: {source_file or 'Unknown'}  
**Extracted**: {timestamp}  
**Tool**: Enhanced TextExtraction with OCR  
**Content Sections**: {len(sections)} total ({text_sections} text, {image_sections} OCR)

---

"""
            
            markdown_content = header + text
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"✅ Enhanced markdown saved: {output_path}")
            print(f"📊 Content extracted: {len(text)} characters, {text_sections} text sections, {image_sections} OCR sections")
            return True
            
        except Exception as e:
            print(f"❌ Failed to save markdown: {e}")
            return False


# Update your main function
def main():
    """Test enhanced extraction with OCR"""
    test_files = [
        # Add more test files here
    ]
    
    # Create output directory
    output_dir = "extracted_markdown_enhanced"
    current_dir = os.getcwd()
    output_folder = os.path.join(current_dir, output_dir)
    os.makedirs(output_folder, exist_ok=True)
    
    print("🚀 Starting Enhanced Text Extraction with OCR...")
    print("=" * 60)
    
    extraction_results = []
    
    for file_path in test_files:
        if not os.path.isfile(file_path):
            print(f"⚠️ File not found: {file_path}")
            continue
            
        # Extract text with OCR
        text = EnhancedTextExtraction.extract_text_from_any(file_path)
        
        if text.startswith("Error") or text.startswith("Warning"):
            print(f"❌ {text}")
            extraction_results.append({
                'file': file_path,
                'status': 'failed',
                'message': text
            })
        else:
            # Show preview
            preview = text[:800] + ("..." if len(text) > 800 else "")
            print(f"📄 Preview:\n{preview}\n")
            
            # Save as enhanced markdown
            base_name = os.path.basename(file_path)
            file_name_without_ext = os.path.splitext(base_name)[0]
            md_filename = f"{file_name_without_ext}_enhanced_extracted.md"
            md_path = os.path.join(output_folder, md_filename)
            
            success = EnhancedTextExtraction.save_as_markdown(text, md_path, base_name)
            
            extraction_results.append({
                'file': file_path,
                'status': 'success' if success else 'failed',
                'output': md_path if success else None,
                'text_length': len(text),
                'has_ocr': 'OCR' in text or 'Image' in text
            })
    
    # Print enhanced summary
    print(f"\n{'=' * 60}")
    print("📊 ENHANCED EXTRACTION SUMMARY")
    print(f"{'=' * 60}")
    
    successful = [r for r in extraction_results if r['status'] == 'success']
    failed = [r for r in extraction_results if r['status'] == 'failed']
    ocr_files = [r for r in successful if r.get('has_ocr', False)]
    
    print(f"📁 Total files processed: {len(extraction_results)}")
    print(f"✅ Successful extractions: {len(successful)}")
    print(f"🖼️ Files with OCR content: {len(ocr_files)}")
    print(f"❌ Failed extractions: {len(failed)}")
    
    if successful:
        total_chars = sum(r['text_length'] for r in successful if 'text_length' in r)
        print(f"📝 Total characters extracted: {total_chars:,}")
        print(f"📂 Output directory: {output_folder}")


if __name__ == "__main__":
    main()
