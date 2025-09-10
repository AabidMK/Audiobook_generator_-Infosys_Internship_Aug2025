import os
import docx
import fitz
import google.generativeai as genai
from datetime import datetime

# -------- CONFIGURE GEMINI --------
genai.configure(api_key="AIzaSyCT_rCoU0H8ob9KRM4cicDNVC1UMD4gm-g")  


# -------- FILE EXTRACTORS --------
def extract_text_from_pdf(file_path):
    text = ""
    num_pages = 0
    with fitz.open(file_path) as doc:
        num_pages = len(doc)
        for page in doc:
            text += page.get_text()
    return text, num_pages


def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text, len(doc.paragraphs)


def extract_text_from_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    num_lines = text.count("\n") + 1
    return text, num_lines


# -------- GEMINI HELPERS --------
def safe_extract(response):
    """Extract text safely from Gemini response"""
    if hasattr(response, "text") and response.text:
        return response.text.strip()
    elif response.candidates and response.candidates[0].content.parts:
        return response.candidates[0].content.parts[0].text.strip()
    else:
        return "⚠️ No valid response text received from Gemini."


def rewrite_text_with_gemini(text):
    model = genai.GenerativeModel("gemini-1.5-flash")
    prompt = f"""Rewrite and enrich the following text for better narration and listener experience.
Focus on making it clear, engaging, and smooth. Keep original meaning:

{text}
"""
    response = model.generate_content(prompt)
    return safe_extract(response)


# -------- CHUNKING --------
def chunk_text(text, max_words=800):
    """Split text into manageable chunks"""
    words = text.split()
    for i in range(0, len(words), max_words):
        yield " ".join(words[i:i + max_words])


# -------- MAIN PROCESS --------
def process_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        raw_text, meta_info = extract_text_from_pdf(file_path)
        meta_label = f"Pages Extracted: {meta_info}"
    elif ext == ".docx":
        raw_text, meta_info = extract_text_from_docx(file_path)
        meta_label = f"Paragraphs Extracted: {meta_info}"
    elif ext == ".txt":
        raw_text, meta_info = extract_text_from_txt(file_path)
        meta_label = f"Lines Extracted: {meta_info}"
    else:
        raise ValueError("Unsupported file format. Use PDF, DOCX, or TXT.")

    word_count = len(raw_text.split())
    char_count = len(raw_text)

    print(f"✅ Extracted {word_count} words. Splitting into chunks...")

    rewritten_chunks = []
    for idx, chunk in enumerate(chunk_text(raw_text, max_words=800)):
        print(f"➡️ Processing chunk {idx+1}...")
        rewritten = rewrite_text_with_gemini(chunk)
        rewritten_chunks.append(rewritten)

    final_text = "\n\n".join(rewritten_chunks)

    # Save outputs into Markdown file
    output_file = f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
    with open(output_file, "w", encoding="utf-8") as f:
        f.write("# 📖 Extraction Summary\n\n")
        f.write(f"- File: {os.path.basename(file_path)}\n")
        f.write(f"- {meta_label}\n")
        f.write(f"- Words Extracted: {word_count}\n")
        f.write(f"- Characters Extracted: {char_count}\n\n")

        f.write("# ✍️ Rewritten & Enriched Text\n\n")
        f.write(final_text + "\n")

    print(f"✅ Rewritten text & summary saved in: {output_file}")


# -------- RUN SCRIPT --------
if __name__ == "__main__":
    file_path = r"c:/Users/admin/Downloads/sampletext.pdf"  # change to your file
    process_file(file_path)
