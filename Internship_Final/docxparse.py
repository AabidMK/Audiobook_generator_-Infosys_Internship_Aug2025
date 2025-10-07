import docx
def extract_docx_text(docx_path):
    doc = docx.Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            text += "\n" + " | ".join(row_data)
    return text