import os
import fitz  
import docx
from datetime import datetime
from PIL import Image
import pytesseract
import easyocr


class TextExtractor:
    """
    Universal text extractor:
    - PDF (text + OCR for scanned pages)
    - DOCX (paragraphs + tables)
    - Images (Tesseract or EasyOCR)
    """

    def __init__(self, use_easyocr=False):
        self.use_easyocr = use_easyocr
        if use_easyocr:
            self.reader = easyocr.Reader(['en'])

    def extract_from_image(self, image_path):
        if not os.path.isfile(image_path):
            return "Error: Image file not found"
        try:
            if self.use_easyocr:
                result = self.reader.readtext(image_path)
                text = "\n".join([res[1] for res in result])
            else:
                img = Image.open(image_path)
                if img.mode != "RGB":
                    img = img.convert("RGB")
                text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
            cleaned = "\n".join(line.strip() for line in text.splitlines() if line.strip())
            return cleaned or "Warning: No text detected in image"
        except Exception as e:
            return f"Error extracting text from image: {e}"

    def extract_from_pdf(self, pdf_path):
        if not os.path.isfile(pdf_path):
            return "Error: PDF file not found"
        try:
            text = []
            with fitz.open(pdf_path) as doc:
                for i, page in enumerate(doc, start=1):
                    page_text = page.get_text().strip()
                    if page_text:
                        text.append(f"## Page {i}\n\n{page_text}")
                    else:
                        pix = page.get_pixmap(dpi=300)
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        if self.use_easyocr:
                            ocr = self.reader.readtext(img)
                            ocr_text = "\n".join([res[1] for res in ocr])
                        else:
                            ocr_text = pytesseract.image_to_string(img, config="--oem 3 --psm 6")
                        text.append(f"## Page {i} (OCR)\n\n{ocr_text.strip()}")
            return "\n\n".join(text) or "Warning: PDF contains no extractable text"
        except Exception as e:
            return f"Error extracting PDF text: {e}"

    def extract_from_docx(self, docx_path):
        if not os.path.isfile(docx_path):
            return "Error: DOCX file not found"
        try:
            doc_obj = docx.Document(docx_path)
            texts = [p.text for p in doc_obj.paragraphs if p.text.strip()]
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            texts.append(cell.text)
            return "\n".join(texts) or "Warning: DOCX contains no extractable text"
        except Exception as e:
            return f"Error extracting DOCX text: {e}"

    def extract_from_any(self, file_path):
        if not os.path.isfile(file_path):
            return f"Error: File '{file_path}' not found"

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_from_pdf(file_path)
        elif ext == ".docx":
            return self.extract_from_docx(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".gif"]:
            return self.extract_from_image(file_path)
        else:
            return "Error: Unsupported file format"

    def save_as_markdown(self, text, output_path, source_file=None):
        try:
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            header = f"""# Extracted Text

**Source**: {source_file or 'Unknown'}  
**Extracted**: {timestamp}  
**Tool**: TextExtractor  

---

"""
            content = header + text
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(content)
            print(f"✅ Saved: {output_path}")
            return True
        except Exception as e:
            print(f"❌ Failed to save {output_path}: {e}")
            return False


def main():
    test_files = [
        # Only PDF, DOCX, and image
        r"c:/Users/admin/Downloads/sampletext.pdf",
        r"C:/Users/admin/Downloads/sampletext.docx",
        r"C:/Users/admin/Downloads/sampleimage.jpg",
    ]

    extractor = TextExtractor(use_easyocr=False)  
    output_dir = "extracted_markdown"
    os.makedirs(output_dir, exist_ok=True)

    results = []
    for file in test_files:
        print(f"\n=== Extracting from: {file} ===")
        text = extractor.extract_from_any(file)
        if text.startswith("Error") or text.startswith("Warning"):
            print(text)
            results.append({"file": file, "status": "failed", "msg": text})
        else:
            preview = text[:300].replace("\n", " ") + ("..." if len(text) > 300 else "")
            print("Preview:", preview)
            base = os.path.splitext(os.path.basename(file))[0]
            md_path = os.path.join(output_dir, f"{base}_extracted.md")
            success = extractor.save_as_markdown(text, md_path, file)
            results.append({
                "file": file,
                "status": "success" if success else "failed",
                "output": md_path if success else None,
                "length": len(text)
            })

   
if __name__ == "__main__":
    main()
