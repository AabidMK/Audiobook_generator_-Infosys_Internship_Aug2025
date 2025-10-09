import fitz  # PyMuPDF
import docx
import os
import signal
from contextlib import contextmanager
from datetime import datetime
import pytesseract
from PIL import Image
import io
import tempfile

@contextmanager
def timeout(duration):
    def timeout_handler(signum, frame):
        raise TimeoutError(f"Operation timed out after {duration} seconds")
    
    # Only use signal on Unix systems
    if hasattr(signal, 'SIGALRM'):
        signal.signal(signal.SIGALRM, timeout_handler)
        signal.alarm(duration)
    try:
        yield
    finally:
        if hasattr(signal, 'SIGALRM'):
            signal.alarm(0)

class EnhancedTextExtraction:
    """Enhanced universal text extraction with OCR support and timeout protection"""
    
    @staticmethod
    def configure_tesseract():
        """Configure Tesseract OCR settings for optimal performance"""
        try:
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!? '
            return custom_config
        except:
            return r'--oem 3 --psm 6'
    
    @staticmethod
    def extract_text_from_image(image_path_or_bytes, is_bytes=False):
        """Enhanced OCR extraction with error handling"""
        supported_formats = ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']
        
        try:
            # Check if tesseract is available
            try:
                import pytesseract
                custom_config = EnhancedTextExtraction.configure_tesseract()
            except ImportError:
                return "OCR not available - pytesseract not installed"
            except:
                return "OCR configuration error"
            
            if is_bytes:
                img = Image.open(io.BytesIO(image_path_or_bytes))
            else:
                if not os.path.isfile(image_path_or_bytes):
                    return "Error: Image file not found"
                
                ext = os.path.splitext(image_path_or_bytes)[1].lower()
                if ext not in supported_formats:
                    return f"Error: Unsupported image format {ext}"
                
                img = Image.open(image_path_or_bytes)
            
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            # Resize small images for better OCR
            width, height = img.size
            if width < 300 or height < 300:
                scale_factor = max(300/width, 300/height)
                new_width = int(width * scale_factor)
                new_height = int(height * scale_factor)
                img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            text = pytesseract.image_to_string(img, config=custom_config)
            cleaned_text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
            
            return cleaned_text or ""
            
        except Exception as e:
            return f"Error extracting text from image: {e}"
    
    @staticmethod
    def extract_text_from_pdf_with_timeout(file_path: str, timeout_per_page: int = 15) -> str:
        """Extract text from PDF with per-page timeout"""
        
        try:
            doc = fitz.open(file_path)
            full_text = []
            total_pages = doc.page_count
            
            print(f"Processing PDF: {total_pages} pages")
            
            for page_num in range(total_pages):
                try:
                    print(f"Processing PDF page {page_num + 1}/{total_pages}...", end=" ", flush=True)
                    
                    # Use timeout only on Unix systems
                    if hasattr(signal, 'SIGALRM'):
                        with timeout(timeout_per_page):
                            page = doc[page_num]
                            text = page.get_text()
                    else:
                        # Windows fallback - no timeout but with progress
                        page = doc[page_num]
                        text = page.get_text()
                    
                    if text.strip():
                        full_text.append(f"--- Page {page_num + 1} Text ---")
                        full_text.append(text)
                        print("SUCCESS")
                    else:
                        print("(empty)")
                    
                    # Process images if available
                    try:
                        image_list = page.get_images(full=True)
                        for img_index, img in enumerate(image_list):
                            try:
                                xref = img[0]
                                pix = fitz.Pixmap(doc, xref)
                                
                                if pix.n - pix.alpha < 4:  # GRAY or RGB
                                    img_data = pix.tobytes("png")
                                    ocr_text = EnhancedTextExtraction.extract_text_from_image(img_data, is_bytes=True)
                                    
                                    if ocr_text and not ocr_text.startswith("Error") and len(ocr_text.strip()) > 10:
                                        full_text.append(f"--- Page {page_num + 1} Image {img_index + 1} OCR ---")
                                        full_text.append(ocr_text)
                                
                                pix = None  # Release memory
                                
                            except Exception as img_error:
                                continue  # Skip problematic images
                    except:
                        pass  # Skip image processing if it fails
                    
                except TimeoutError:
                    print(f"⏰ TIMEOUT (skipped)")
                    continue
                    
                except Exception as e:
                    print(f"ERROR: {e}")
                    continue
                
                # Memory cleanup every 10 pages
                if page_num % 10 == 0 and page_num > 0:
                    print(f"Memory cleanup at page {page_num}")
                    import gc
                    gc.collect()
                
                # Safety break for very large documents
                if page_num > 100:
                    print("Limiting to first 100 pages for safety")
                    break
            
            doc.close()
            
            final_text = '\n\n'.join(full_text)
            print(f"Extraction complete: {len(final_text):,} characters")
            
            return final_text if final_text else "No text extracted"
            
        except Exception as e:
            return f"Error extracting from PDF: {e}"
    
    @staticmethod
    def extract_text_from_pdf_basic(file_path: str) -> str:
        """Basic PDF extraction as fallback"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = []
                
                total_pages = len(reader.pages)
                print(f"Fallback extraction: {total_pages} pages")
                
                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text.append(page_text)
                        print(f"Page {i+1}/{total_pages} SUCCESS")
                        
                        # Safety limit
                        if i > 50:
                            print("Limiting to first 50 pages")
                            break
                            
                    except Exception as e:
                        print(f"Page {i+1} error: {e}")
                        continue
                
                return '\n\n'.join(text) if text else "No text extracted"
                
        except Exception as e:
            return f"Error with basic PDF extraction: {e}"
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Enhanced DOCX extraction"""
        try:
            doc = docx.Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n\n'.join(text_parts) if text_parts else "No text found in DOCX"
            
        except Exception as e:
            return f"Error extracting DOCX: {e}"
    
    @staticmethod
    def extract_text_from_any(file_path: str) -> str:
        """Universal extractor with timeout protection"""
        
        if not os.path.exists(file_path):
            return f"Error: File not found - {file_path}"
        
        file_extension = os.path.splitext(file_path)[1].lower()
        image_formats = ['.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif']
        
        print(f"Processing: {os.path.basename(file_path)}")
        
        try:
            if file_extension == '.pdf':
                # Try enhanced PDF extraction first
                result = EnhancedTextExtraction.extract_text_from_pdf_with_timeout(file_path, timeout_per_page=15)
                
                if not result.startswith("Error") and len(result) > 100:
                    return result
                
                # Fallback to basic extraction
                print("🔄 Trying fallback PDF extraction...")
                return EnhancedTextExtraction.extract_text_from_pdf_basic(file_path)
            
            elif file_extension == '.docx':
                return EnhancedTextExtraction.extract_text_from_docx(file_path)
            
            elif file_extension == '.txt':
                encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            content = f.read()
                            return content if content.strip() else "Empty text file"
                    except UnicodeDecodeError:
                        continue
                return "Error: Unable to decode text file"
            
            elif file_extension == '.md':
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        return content if content.strip() else "Empty markdown file"
                except Exception as e:
                    return f"Error reading markdown file: {e}"
            
            elif file_extension in image_formats:
                return EnhancedTextExtraction.extract_text_from_image(file_path)
            
            else:
                # Try as text file
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        return content if content.strip() else "Empty file"
                except UnicodeDecodeError:
                    return f"Error: Unsupported file type - {file_extension}"
                except Exception as e:
                    return f"Error reading file: {e}"
                    
        except Exception as e:
            return f"Error during extraction: {e}"

# Test function
def test_extraction():
    """Test the extraction with a sample file"""
    test_file = "testing.pdf"  # Change this to your file
    
    if os.path.exists(test_file):
        print(f"Testing extraction on: {test_file}")
        result = EnhancedTextExtraction.extract_text_from_any(test_file)
        print(f"\nResult length: {len(result)} characters")
        print(f"First 500 characters:\n{result[:500]}...")
    else:
        print(f"Test file not found: {test_file}")

if __name__ == "__main__":
    test_extraction()
