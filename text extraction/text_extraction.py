import fitz  # PyMuPDF
import docx
import os
from datetime import datetime
import pytesseract
from PIL import Image



class TextExtraction:
    """
    Robust universal text extraction for PDF, DOCX, TXT, images, and most readable plain text files.
    Uses best-in-class parsers including Tesseract OCR for images.
    """

    @staticmethod
    def extract_text_from_image(image_path):
        """
        Extracts text from an image file using Tesseract OCR.
        Supports common formats: PNG, JPG, JPEG, TIFF, BMP, GIF
        Returns extracted text or informative error message.
        """
        if not os.path.isfile(image_path):
            return "Error: Image file not found"
        
        supported_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}
        ext = os.path.splitext(image_path)[1].lower()
        
        if ext not in supported_formats:
            return f"Error: Unsupported image format '{ext}'. Supported: {', '.join(supported_formats)}"
        
        try:
            # Configure Tesseract (uncomment and adjust path if needed)
            # pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'  # macOS Homebrew path
            
            # Open and process image
            with Image.open(image_path) as img:
                # Convert to RGB if necessary (for better OCR results)
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                
                # Extract text using Tesseract with optimized config
                custom_config = r'--oem 3 --psm 6'
                text = pytesseract.image_to_string(img, config=custom_config)
                
                # Clean up the extracted text
                cleaned_text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())
                
                return cleaned_text or "Warning: No text detected in image"
                
        except Exception as e:
            return f"Error extracting text from image: {e}"

    @staticmethod
    def extract_text_from_pdf(pdf_path):
        """
        Extracts text from a PDF file using PyMuPDF.
        Returns extracted text or informative error message.
        """
        if not os.path.isfile(pdf_path):
            return "Error: PDF file not found"
        try:
            with fitz.open(pdf_path) as doc:
                text = [page.get_text() for page in doc]
            return '\n'.join(text).strip() or "Warning: PDF contains no extractable text"
        except Exception as e:
            return f"Error extracting PDF text: {e}"


    @staticmethod
    def extract_text_from_docx(docx_path):
        """
        Extracts text from a DOCX file using python-docx, including paragraphs and tables.
        Returns extracted text or informative error message.
        """
        if not os.path.isfile(docx_path):
            return "Error: DOCX file not found"
        try:
            doc_obj = docx.Document(docx_path)
            texts = [para.text for para in doc_obj.paragraphs]
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        texts.append(cell.text)
            result = '\n'.join([line for line in texts if line.strip()])
            return result or "Warning: DOCX contains no extractable text"
        except Exception as e:
            return f"Error extracting DOCX text: {e}"


    @staticmethod
    def extract_text_from_txt(txt_path):
        """
        Extracts text from a TXT file, auto-trying several common encodings.
        Returns extracted text or informative error message.
        """
        if not os.path.isfile(txt_path):
            return "Error: TXT file not found"
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        for encoding in encodings:
            try:
                with open(txt_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return f"Error extracting TXT text: {e}"
        return "Error: Unable to decode text file with UTF-8/16/latin-1/cp1252"


    @staticmethod
    def extract_text_from_any(file_path):
        """
        Universal extractor:
        Uses best parser for .pdf, .docx, .txt, and image extensions.
        For others, attempts UTF-8 plain text extraction as fallback.
        """
        if not os.path.isfile(file_path):
            return f"Error: File '{file_path}' not found"
        
        ext = os.path.splitext(file_path)[1].lower()
        
        # Image formats
        image_formats = {'.png', '.jpg', '.jpeg', '.tiff', '.tif', '.bmp', '.gif'}
        
        if ext == '.pdf':
            return TextExtraction.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return TextExtraction.extract_text_from_docx(file_path)
        elif ext == '.txt':
            return TextExtraction.extract_text_from_txt(file_path)
        elif ext in image_formats:
            return TextExtraction.extract_text_from_image(file_path)
        else:
            # safe attempt to read as UTF-8 text
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                return "Error: Unknown file type and not UTF-8 decodable"
            except Exception as e:
                return f"Error reading unknown file type: {e}"


    @staticmethod
    def save_as_markdown(text, output_path, source_file=None):
        """
        Save extracted text as a markdown (.md) file with metadata header.
        """
        try:
            # Create markdown header with metadata
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            header = f"""# Extracted Text


**Source**: {source_file or 'Unknown'}  
**Extracted**: {timestamp}  
**Tool**: TextExtraction Module


---


"""
            
            # Combine header with content
            markdown_content = header + text
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            print(f"✅ Markdown file saved: {output_path}")
            return True
            
        except Exception as e:
            print(f"❌ Failed to save markdown file {output_path}: {e}")
            return False



def main():
    """Enhanced main function with image support"""
    test_files = [
        #"/Users/adeshnarayanatellakua/Documents/adobe_project/input/1.pdf",
        #"/Users/adeshnarayanatellakua/Downloads/RUST_Programming_Lab_manual.docx", 
        #"/Users/adeshnarayanatellakua/Downloads/long-doc.txt",
        # Add your image files here:
        "/Users/adeshnarayanatellakua/Desktop/Screenshot.png",
        # "/path/to/your/screenshot.jpg",
    ]
    
    # Create output directory in current working directory
    output_dir = "extracted_markdown"
    current_dir = os.getcwd()
    output_folder = os.path.join(current_dir, output_dir)
    
    # Create output folder if it doesn't exist
    os.makedirs(output_folder, exist_ok=True)
    
    extraction_results = []
    
    for file_path in test_files:
        print(f"\n=== Extracting from: {file_path} ===")
        
        # Extract text
        text = TextExtraction.extract_text_from_any(file_path)
        
        if text.startswith("Error") or text.startswith("Warning"):
            print(text)
            extraction_results.append({
                'file': file_path,
                'status': 'failed',
                'message': text
            })
        else:
            # Show preview
            preview = text[:500] + ("..." if len(text) > 500 else "")
            print(preview)
            
            # Generate markdown filename using only the base filename
            base_name = os.path.basename(file_path)
            file_name_without_ext = os.path.splitext(base_name)[0]
            md_filename = f"{file_name_without_ext}_extracted.md"
            
            # Save in project folder
            md_path = os.path.join(output_folder, md_filename)
            
            # Save as markdown
            success = TextExtraction.save_as_markdown(text, md_path, base_name)
            
            extraction_results.append({
                'file': file_path,
                'status': 'success' if success else 'failed',
                'output': md_path if success else None,
                'text_length': len(text)
            })
    
    # Print summary
    print(f"\n{'='*50}")
    print("📊 EXTRACTION SUMMARY")
    print(f"{'='*50}")
    
    successful = [r for r in extraction_results if r['status'] == 'success']
    failed = [r for r in extraction_results if r['status'] == 'failed']
    
    print(f"Total files processed: {len(extraction_results)}")
    print(f"Successful extractions: {len(successful)}")
    print(f"Failed extractions: {len(failed)}")
    
    if successful:
        total_chars = sum(r['text_length'] for r in successful if 'text_length' in r)
        print(f"Total characters extracted: {total_chars:,}")
        print(f"✅ Output directory: {output_folder}")
        
        print("\n📄 Successfully extracted files:")
        for result in successful:
            original_name = os.path.basename(result['file'])
            output_name = os.path.basename(result['output'])
            print(f"  • {original_name} → {output_name}")
    
    if failed:
        print("\n❌ Failed extractions:")
        for result in failed:
            original_name = os.path.basename(result['file'])
            print(f"  • {original_name}: {result['message']}")



if __name__ == "__main__":
    main()
